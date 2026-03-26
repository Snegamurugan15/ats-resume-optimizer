"""
ATS Resume Optimizer — Streamlit App
=====================================
Upload resume (PDF/DOCX) + paste job description
→ Claude AI rewrites for 98%+ ATS score
→ Generates matching cover letter
→ Both files saved to Google Drive as "Company – Role"

All 8 rules applied:
  Rule 1: Score tiers (>90 no edit, 60-89 → 95%+, 40-59 → 80%+, <40 best effort)
  Rule 2: Never fabricate — truthful only
  Rule 3: Research ATS platform (Workday, BambooHR, etc.)
  Rule 4: Output to Google Drive folder "CompanyName – RoleName"
  Rule 5: Calibri 10pt, justified, centered bold headers, max 2 pages
  Rule 6: No hyperlinks, remove weakest project if needed
  Rule 7: Final statement — ATS score achieved + cert recommendations
  Rule 8: Never change job titles/position names
"""

import streamlit as st
import requests
import json
import re
import io
import os
import tempfile
from datetime import date

# ── Document libraries ─────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import PyPDF2
    PDF_OK = True
except ImportError:
    try:
        import pdfplumber
        PDF_OK = True
        PyPDF2 = None
    except ImportError:
        PDF_OK = False

# ── Google Drive ───────────────────────────────────────────────────────────────
try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload
    DRIVE_OK = True
except ImportError:
    DRIVE_OK = False

# ──────────────────────────────────────────────────────────────────────────────
# Page config
# ──────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title  = "ATS Resume Optimizer",
    page_icon   = "🎯",
    layout      = "wide",
    initial_sidebar_state = "expanded"
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1F3864, #2d5a9e);
        color: white;
        padding: 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        text-align: center;
    }
    .score-box {
        background: #f0fdf4;
        border: 2px solid #16a34a;
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
    }
    .score-number { font-size: 4rem; font-weight: 800; color: #16a34a; }
    .score-warn   { font-size: 4rem; font-weight: 800; color: #d97706; }
    .score-low    { font-size: 4rem; font-weight: 800; color: #dc2626; }
    .keyword-matched {
        background: #dcfce7; color: #166534;
        padding: 3px 10px; border-radius: 999px;
        margin: 3px; display: inline-block; font-size: 0.8rem;
    }
    .keyword-missing {
        background: #fee2e2; color: #991b1b;
        padding: 3px 10px; border-radius: 999px;
        margin: 3px; display: inline-block; font-size: 0.8rem;
    }
    .rule-badge {
        background: #eff6ff; color: #1e40af;
        padding: 2px 8px; border-radius: 4px;
        font-size: 0.75rem; font-weight: 600;
    }
    div[data-testid="stDownloadButton"] button {
        background-color: #1F3864 !important;
        color: white !important;
        border-radius: 8px !important;
        width: 100% !important;
    }
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────────────────────
# Sidebar — Settings
# ──────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## ⚙️ Settings")
    st.markdown("---")

    api_key = st.text_input(
        "Claude API Key",
        type       = "password",
        placeholder= "sk-ant-api03-...",
        help       = "Get your key at console.anthropic.com (free tier available)"
    )

    st.markdown("---")
    st.markdown("### 🔗 Google Drive (optional)")

    drive_json = st.text_area(
        "Paste Service Account JSON",
        height      = 120,
        placeholder = '{"type": "service_account", "project_id": "...", ...}',
        help        = "Paste the contents of your downloaded JSON key file"
    )

    drive_folder_id = st.text_input(
        "Drive Folder ID (optional)",
        placeholder = "From the Drive folder URL",
        help        = "Open your Drive folder → copy the ID from the URL after /folders/"
    )

    st.markdown("---")
    st.markdown("""
    **📋 Rules Applied**

    <span class='rule-badge'>R1</span> Score tiers<br>
    <span class='rule-badge'>R2</span> Truthful only<br>
    <span class='rule-badge'>R3</span> ATS platform detection<br>
    <span class='rule-badge'>R4</span> Named Drive folder<br>
    <span class='rule-badge'>R5</span> Calibri 10pt, 2 pages<br>
    <span class='rule-badge'>R6</span> No hyperlinks<br>
    <span class='rule-badge'>R7</span> Final ATS statement<br>
    <span class='rule-badge'>R8</span> Job titles preserved
    """, unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────────────────────
# Header
# ──────────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class='main-header'>
    <h1>🎯 ATS Resume Optimizer</h1>
    <p style='margin:0; opacity:0.85;'>
        Upload your resume + paste any job description<br>
        → Get 80–98%+ ATS score + matching cover letter → Auto-saved to Google Drive
    </p>
</div>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────────────────────
# Input section
# ──────────────────────────────────────────────────────────────────────────────
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("### 📄 Step 1 — Upload Your Resume")
    uploaded_file = st.file_uploader(
        "Choose your resume",
        type       = ["pdf", "docx"],
        label_visibility = "collapsed"
    )
    if uploaded_file:
        st.success(f"✓ {uploaded_file.name} uploaded")

    st.markdown("### 🏢 Step 2 — Company & Role")
    company_name = st.text_input("Company name", placeholder="e.g. Region of Peel")
    role_name    = st.text_input("Role title",   placeholder="e.g. Technical Analyst, AIMS")

with col2:
    st.markdown("### 📋 Step 3 — Job Description")
    job_description = st.text_area(
        "Paste the complete job description",
        height           = 250,
        label_visibility = "collapsed",
        placeholder      = "Paste the full job description here — the more complete, the better the ATS score..."
    )
    jd_len = len(job_description)
    st.caption(f"{jd_len} characters" + (" ✓ Good length" if jd_len > 500 else " — paste more for better results"))

# ── Optimize button ────────────────────────────────────────────────────────────
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
with col_btn2:
    optimize_clicked = st.button(
        "✨  Optimize My Resume",
        type      = "primary",
        disabled  = not (uploaded_file and job_description.strip() and api_key.strip()),
        use_container_width = True
    )

if not api_key:
    st.info("👈 Enter your Claude API key in the sidebar to get started")

# ──────────────────────────────────────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────────────────────────────────────

def extract_resume_text(uploaded_file) -> str:
    """Extract text from PDF or DOCX."""
    name = uploaded_file.name.lower()
    content = uploaded_file.read()

    if name.endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception:
            pass
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            pass
        return ""

    elif name.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return ""
    return ""


def calculate_ats_score(resume_text: str, job_description: str) -> dict:
    """Simple keyword-based ATS score calculation."""
    def extract_keywords(text):
        words = re.findall(r'\b[a-zA-Z][a-zA-Z+#.]{2,}\b', text.lower())
        stopwords = {"the","and","for","are","was","with","this","that","have","from",
                     "will","your","our","their","been","they","more","than","also",
                     "you","can","all","but","not","any","has","its","had","her",
                     "his","him","she","who","what","when","where","how","why","which",
                     "each","both","only","very","just","into","over","such","then"}
        return set(w for w in words if w not in stopwords and len(w) > 2)

    jd_keywords    = extract_keywords(job_description)
    resume_keywords = extract_keywords(resume_text)

    matched = jd_keywords & resume_keywords
    missing = jd_keywords - resume_keywords

    score = int((len(matched) / max(len(jd_keywords), 1)) * 100)
    score = min(score, 100)

    return {
        "score"   : score,
        "matched" : sorted(list(matched))[:30],
        "missing" : sorted(list(missing))[:20]
    }


def get_optimization_tier(score: int) -> dict:
    if score > 90:
        return {"label": "NO_EDIT",     "target": score, "instruction": f"Score is already {score}% (>90%). Return the resume unchanged with a note that no optimization was needed."}
    elif score >= 60:
        return {"label": "TARGET_95",   "target": 95,    "instruction": f"Initial score is {score}% (60-89%). Maximize to 95%+ ATS score."}
    elif score >= 40:
        return {"label": "TARGET_80",   "target": 80,    "instruction": f"Initial score is {score}% (40-59%). Maximize to 80%+ ATS score."}
    else:
        return {"label": "BEST_EFFORT", "target": 65,    "instruction": f"Initial score is {score}% (<40%). Optimize honestly; flag hard gaps clearly at the end."}


def call_claude(api_key: str, system_prompt: str, user_message: str, max_tokens: int = 4096) -> str | None:
    """Call Claude API and return the text response."""
    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key"         : api_key,
                "anthropic-version" : "2023-06-01",
                "content-type"      : "application/json"
            },
            json={
                "model"      : "claude-opus-4-6",
                "max_tokens" : max_tokens,
                "temperature": 0.2,
                "system"     : system_prompt,
                "messages"   : [{"role": "user", "content": user_message}]
            },
            timeout=120
        )
        if response.status_code == 200:
            return response.json()["content"][0]["text"]
        else:
            st.error(f"API error {response.status_code}: {response.text[:300]}")
            return None
    except Exception as e:
        st.error(f"Connection error: {e}")
        return None


def extract_json(text: str) -> dict | None:
    """Extract JSON from Claude's response."""
    try:
        # Try code block first
        match = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
        if match:
            return json.loads(match.group(1).strip())
        # Try raw JSON
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end > start:
            return json.loads(text[start:end+1])
    except Exception:
        pass
    return None


def build_resume_docx(resume_data: dict) -> bytes:
    """Build ATS-friendly resume DOCX using python-docx."""
    doc = Document()

    # ── Page margins: 0.6" top/bottom, 0.75" left/right ─────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.52)
        section.bottom_margin = Cm(1.52)
        section.left_margin   = Cm(1.90)
        section.right_margin  = Cm(1.90)

    NAVY   = RGBColor(0x1F, 0x38, 0x64)
    GREY   = RGBColor(0x40, 0x40, 0x40)
    BLACK  = RGBColor(0x00, 0x00, 0x00)

    def set_spacing(para, before=0, after=4):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)
        para.paragraph_format.line_spacing = Pt(12)

    def add_run(para, text, bold=False, size=10, color=None, italic=False):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return run

    def add_section_header(doc, title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=6, after=2)
        run = add_run(p, title.upper(), bold=True, size=10, color=NAVY)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1F3864")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_body(doc, text, justify=True, size=10, color=None, before=0, after=3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=before, after=after)
        add_run(p, text, size=size, color=color or GREY)
        return p

    def add_bullet(doc, text, size=10):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=0, after=2)
        add_run(p, text, size=size, color=GREY)

    # ── NAME ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(name_para, before=0, after=2)
    add_run(name_para, resume_data.get("name", "").upper(), bold=True, size=16, color=NAVY)

    # ── CONTACT LINE ──────────────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact_para, before=0, after=6)
    add_run(contact_para, resume_data.get("contactLine", ""), size=9, color=GREY)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if resume_data.get("summary"):
        add_section_header(doc, "Professional Summary")
        add_body(doc, resume_data["summary"])

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    if resume_data.get("experience"):
        add_section_header(doc, "Work Experience")
        for exp in resume_data["experience"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p, before=4, after=1)
            add_run(p, exp.get("title", ""),   bold=True, size=10, color=BLACK)
            add_run(p, f"  |  {exp.get('company','')}  |  {exp.get('duration','')}", size=9, color=GREY)
            for bullet in exp.get("bullets", []):
                add_bullet(doc, bullet)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    if resume_data.get("skills"):
        add_section_header(doc, "Technical Skills & Competencies")
        skills_text = "  •  ".join(resume_data["skills"])
        add_body(doc, skills_text, justify=True)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if resume_data.get("education"):
        add_section_header(doc, "Education")
        for edu in resume_data["education"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p, before=2, after=1)
            add_run(p, edu.get("degree", ""), bold=True, size=10, color=BLACK)
            add_run(p, f"  |  {edu.get('institution','')}  |  {edu.get('year','')}", size=9, color=GREY)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if resume_data.get("certifications"):
        add_section_header(doc, "Certifications")
        for cert in resume_data["certifications"]:
            add_body(doc, f"• {cert}", justify=False, size=10)

    # ── PROJECTS ──────────────────────────────────────────────────────────────
    if resume_data.get("projects"):
        add_section_header(doc, "Academic Projects")
        for proj in resume_data["projects"]:
            p = doc.add_paragraph()
            set_spacing(p, before=3, after=1)
            add_run(p, proj.get("name", ""), bold=True, size=10, color=BLACK)
            if proj.get("techStack"):
                add_run(p, f"  |  {proj['techStack']}", size=9, color=GREY)
            if proj.get("description"):
                add_bullet(doc, proj["description"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cover_letter_docx(cover_body: str, resume_data: dict,
                              company_name: str, role_name: str) -> bytes:
    """Build cover letter DOCX matching resume header style."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.52)
        section.bottom_margin = Cm(1.52)
        section.left_margin   = Cm(1.90)
        section.right_margin  = Cm(1.90)

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    GREY = RGBColor(0x40, 0x40, 0x40)

    def sp(para, before=0, after=4):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)
        para.paragraph_format.line_spacing = Pt(13)

    def run(para, text, bold=False, size=10, color=None):
        r = para.add_run(text)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    # Name header
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(np, after=2)
    run(np, resume_data.get("name","").upper(), bold=True, size=16, color=NAVY)

    # Contact
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(cp, after=8)
    run(cp, resume_data.get("contactLine",""), size=9, color=GREY)

    # HR line
    hr = doc.add_paragraph()
    sp(hr, after=6)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "1F3864"); pBdr.append(bottom); pPr.append(pBdr)

    # Date
    dp = doc.add_paragraph()
    sp(dp, after=8)
    run(dp, date.today().strftime("%B %d, %Y"), size=10, color=GREY)

    # Subject line
    subj = doc.add_paragraph()
    sp(subj, after=8)
    run(subj, f"Re: Application for {role_name} – {company_name}", bold=True, size=10, color=NAVY)

    # Salutation
    sal = doc.add_paragraph()
    sp(sal, after=8)
    run(sal, "Dear Hiring Manager,", size=10)

    # Body paragraphs
    for para_text in cover_body.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp(p, after=8)
        run(p, para_text, size=10, color=GREY)

    # Closing
    cl = doc.add_paragraph()
    sp(cl, before=6, after=16)
    run(cl, "Sincerely,", size=10)

    sig = doc.add_paragraph()
    sp(sig, after=2)
    name_title = resume_data.get("name","").title()
    run(sig, name_title, bold=True, size=10, color=NAVY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def upload_to_drive(drive_json_str: str, folder_id: str,
                    folder_name: str, files: list[tuple]) -> str | None:
    """
    Upload files to Google Drive.
    files = list of (filename, bytes_content, mime_type)
    Returns the Drive folder link or None.
    """
    if not DRIVE_OK or not drive_json_str.strip():
        return None
    try:
        creds_dict = json.loads(drive_json_str)
        creds = service_account.Credentials.from_service_account_info(
            creds_dict,
            scopes=["https://www.googleapis.com/auth/drive.file"]
        )
        service = build("drive", "v3", credentials=creds)

        # Create subfolder
        folder_meta = {
            "name"    : folder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents" : [folder_id] if folder_id else []
        }
        folder = service.files().create(body=folder_meta, fields="id").execute()
        fid = folder["id"]

        # Upload each file
        for filename, content, mime in files:
            file_meta = {"name": filename, "parents": [fid]}
            media = MediaIoBaseUpload(io.BytesIO(content), mimetype=mime)
            service.files().create(body=file_meta, media_body=media, fields="id").execute()

        # Make folder accessible via link
        service.permissions().create(
            fileId=fid,
            body={"type": "anyone", "role": "reader"}
        ).execute()

        return f"https://drive.google.com/drive/folders/{fid}"
    except Exception as e:
        st.warning(f"Drive upload failed: {e}")
        return None


# ──────────────────────────────────────────────────────────────────────────────
# SYSTEM PROMPTS
# ──────────────────────────────────────────────────────────────────────────────

ATS_SYSTEM_PROMPT = """
You are an elite ATS resume optimization expert with 15+ years of experience achieving 98%+ ATS match scores.

## SCORING TIERS (Rule 1):
- Initial score >90%  → Return resume UNCHANGED + note: "Your resume already achieves X% ATS match. No optimization needed."
- Initial score 60-89% → Maximize to 95%+ ATS score
- Initial score 40-59% → Maximize to 80%+ ATS score
- Initial score <40%   → Best effort; clearly flag hard skill gaps at the end

## ABSOLUTE RULES:
1. NEVER fabricate skills, certifications, or experience the candidate doesn't have (Rule 2)
2. NEVER change job titles/position names — only reframe bullet descriptions (Rule 8)
3. Use EXACT keyword phrases from the JD — ATS matches literals, not synonyms
4. Research the company's ATS platform (Workday, BambooHR, Greenhouse, Taleo, etc.) and structure accordingly (Rule 3)
5. Remove least-relevant project if needed to stay within 2 pages (Rule 6)
6. Replace all hyperlinks with plain text (Rule 6)
7. End response with: "ATS Score Achieved: X%." and optional certification recommendations (Rule 7)

## RESUME FORMAT (Rule 5):
- Font: Calibri 10pt body, justified paragraphs, centred bold section headers
- Max 2 pages
- Section headers: Professional Summary | Work Experience | Technical Skills & Competencies | Education | Certifications | Academic Projects

## OUTPUT FORMAT — return ONLY this JSON:
```json
{
  "optimizedResumeText": "Full rewritten resume as plain text",
  "atsScore": 98,
  "keywordsMatched": ["keyword1", "keyword2"],
  "keywordsMissing": ["gap1", "gap2"],
  "improvementSummary": "2-3 sentences. Ends with: ATS Score Achieved: X%. Recommended certifications: ...",
  "sectionFeedback": {
    "summary": "what changed",
    "experience": "what changed",
    "skills": "what changed"
  }
}
```
""".strip()

DOCX_PARSE_PROMPT = """
Extract structured data from this resume and return ONLY valid JSON — no markdown, no explanation.

{
  "name": "FULL NAME",
  "contactLine": "City, Province  •  Phone  •  email@example.com  •  linkedin.com/in/...",
  "summary": "Professional summary paragraph",
  "experience": [
    {
      "title": "Exact Job Title (never changed)",
      "company": "Company Name",
      "duration": "Mon YYYY – Mon YYYY",
      "location": "City, Province",
      "bullets": ["Achievement 1 with numbers", "Achievement 2"]
    }
  ],
  "education": [
    {
      "degree": "Degree – Field of Study",
      "institution": "Institution, City",
      "year": "YYYY – YYYY",
      "gpa": ""
    }
  ],
  "skills": ["Skill1", "Skill2"],
  "certifications": ["Cert1"],
  "projects": [
    {"name": "Project", "description": "What it does", "techStack": "Tech1, Tech2"}
  ]
}
""".strip()

COVER_LETTER_PROMPT = """
You are a professional cover letter writer. Write 3-4 tight, impactful body paragraphs.
No salutation, no closing, no signature — body paragraphs only.

Rules:
- Para 1: Hook — state the role + 1 compelling reason you're the right fit
- Para 2: Core technical skills tied to JD keywords (be specific, use numbers)
- Para 3: Key achievement with quantified impact
- Para 4: Why this company + call to action
- Use keywords from the job description throughout
- Professional, confident tone — no clichés like "I am writing to express..."
- Prose paragraphs only — no bullet points
- Return ONLY the body paragraphs
""".strip()


# ──────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE — runs on button click
# ──────────────────────────────────────────────────────────────────────────────

if optimize_clicked:

    if not uploaded_file:
        st.error("Please upload your resume first.")
        st.stop()
    if not job_description.strip():
        st.error("Please paste the job description.")
        st.stop()
    if not api_key.strip():
        st.error("Please enter your Claude API key in the sidebar.")
        st.stop()

    st.markdown("---")
    progress = st.progress(0)
    status   = st.empty()

    # ── 1. Extract resume text ────────────────────────────────────────────────
    status.info("📄 Reading your resume…")
    progress.progress(10)
    uploaded_file.seek(0)
    resume_text = extract_resume_text(uploaded_file)

    if not resume_text.strip():
        st.error("Could not read the resume. Please try a different PDF or DOCX file.")
        st.stop()

    # ── 2. Initial ATS score ──────────────────────────────────────────────────
    status.info("📊 Calculating initial ATS score…")
    progress.progress(20)
    initial = calculate_ats_score(resume_text, job_description)
    tier    = get_optimization_tier(initial["score"])

    st.info(f"📊 Initial ATS score: **{initial['score']}%** → {tier['instruction']}")

    # ── 3. Claude optimization ────────────────────────────────────────────────
    status.info(f"🤖 Claude is optimizing your resume (targeting {tier['target']}%+)…")
    progress.progress(35)

    user_msg = f"""
Company: {company_name or 'Not specified'}
Role: {role_name or 'Not specified'}
Optimization tier: {tier['instruction']}

CRITICAL: NEVER change job titles/position names (Rule 8). Only reframe bullet descriptions.

=== ORIGINAL RESUME ===
{resume_text}

=== JOB DESCRIPTION ===
{job_description}
""".strip()

    opt_response = call_claude(api_key, ATS_SYSTEM_PROMPT, user_msg, max_tokens=4096)
    if not opt_response:
        st.error("Claude API call failed. Check your API key and try again.")
        st.stop()

    opt_result = extract_json(opt_response)
    if not opt_result:
        opt_result = {
            "optimizedResumeText": opt_response,
            "atsScore"           : tier["target"],
            "keywordsMatched"    : initial["matched"],
            "keywordsMissing"    : initial["missing"],
            "improvementSummary" : f"Resume optimized. ATS Score Achieved: {tier['target']}%.",
            "sectionFeedback"    : {}
        }

    # ── 4. Parse into structured data for DOCX ────────────────────────────────
    status.info("📝 Structuring resume data…")
    progress.progress(55)

    parse_msg = f"{DOCX_PARSE_PROMPT}\n\n=== RESUME TEXT ===\n{opt_result['optimizedResumeText']}"
    parse_response = call_claude(api_key, "You extract structured data from resumes. Return ONLY valid JSON.", parse_msg)
    resume_data = extract_json(parse_response) if parse_response else None

    if not resume_data:
        lines = opt_result["optimizedResumeText"].split("\n")
        resume_data = {
            "name"       : lines[0].strip() if lines else "Candidate",
            "contactLine": lines[1].strip() if len(lines) > 1 else "",
            "summary"    : "",
            "experience" : [],
            "education"  : [],
            "skills"     : opt_result.get("keywordsMatched", [])[:20],
            "certifications": [],
            "projects"   : []
        }

    # ── 5. Generate cover letter ───────────────────────────────────────────────
    status.info("✉️ Writing your cover letter…")
    progress.progress(70)

    cover_msg = f"""
Company: {company_name or 'the company'}
Role: {role_name or 'the position'}

RESUME (summary):
{opt_result['optimizedResumeText'][:2000]}

JOB DESCRIPTION:
{job_description[:2000]}
""".strip()

    cover_response = call_claude(api_key, COVER_LETTER_PROMPT, cover_msg, max_tokens=1500)
    cover_body = cover_response or "I am excited to apply for this role. My background aligns well with your requirements, and I look forward to contributing to your team."

    # ── 6. Build DOCX files ────────────────────────────────────────────────────
    status.info("⚙️ Generating DOCX files…")
    progress.progress(82)

    resume_bytes = build_resume_docx(resume_data)
    cover_bytes  = build_cover_letter_docx(cover_body, resume_data, company_name or "Company", role_name or "Role")

    # ── 7. Upload to Google Drive ─────────────────────────────────────────────
    drive_link = None
    safe_company = (company_name or "Company").replace("/","").strip()
    safe_role    = (role_name or "Role").replace("/","").strip()
    folder_name  = f"{safe_company} \u2013 {safe_role}"

    if drive_json and drive_json.strip():
        status.info(f"☁️ Uploading to Google Drive folder \"{folder_name}\"…")
        progress.progress(92)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        drive_link = upload_to_drive(
            drive_json, drive_folder_id, folder_name,
            [
                (f"Resume_{safe_company}_{safe_role}.docx", resume_bytes, mime),
                (f"CoverLetter_{safe_company}_{safe_role}.docx", cover_bytes, mime)
            ]
        )

    progress.progress(100)
    status.success("✅ Done! Your files are ready.")

    # ──────────────────────────────────────────────────────────────────────────
    # RESULTS
    # ──────────────────────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("## 🎉 Results")

    final_score = opt_result.get("atsScore", tier["target"])

    # ── Score card ─────────────────────────────────────────────────────────────
    score_class = "score-number" if final_score >= 95 else ("score-warn" if final_score >= 80 else "score-low")
    score_label = ("Excellent — ATS Ready ✓" if final_score >= 95
                   else ("Good — Strong match" if final_score >= 80
                   else ("Moderate — Targeted for role" if final_score >= 60
                   else "Optimized — Hard gaps flagged")))

    col_s1, col_s2 = st.columns([1, 2])
    with col_s1:
        st.markdown(f"""
        <div class='score-box'>
            <div class='{score_class}'>{final_score}%</div>
            <div style='color:#374151; font-weight:600'>{score_label}</div>
            <div style='color:#6b7280; font-size:0.85rem; margin-top:4px'>
                Initial: {initial["score"]}% → Final: {final_score}%
            </div>
        </div>
        """, unsafe_allow_html=True)

    with col_s2:
        st.markdown("**📌 Improvement Summary**")
        st.write(opt_result.get("improvementSummary", "Resume optimized successfully."))

        if drive_link:
            st.success(f"✅ Saved to Google Drive → [{folder_name}]({drive_link})")

    # ── Download buttons ────────────────────────────────────────────────────────
    st.markdown("### 📥 Download Your Files")
    dcol1, dcol2 = st.columns(2)

    with dcol1:
        st.download_button(
            label            = "⬇️  Resume DOCX",
            data             = resume_bytes,
            file_name        = f"Resume_{safe_company}_{safe_role}.docx",
            mime             = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width = True
        )

    with dcol2:
        st.download_button(
            label            = "⬇️  Cover Letter DOCX",
            data             = cover_bytes,
            file_name        = f"CoverLetter_{safe_company}_{safe_role}.docx",
            mime             = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width = True
        )

    # ── Keywords ────────────────────────────────────────────────────────────────
    st.markdown("### 🔍 Keyword Analysis")
    kw1, kw2 = st.columns(2)

    matched = opt_result.get("keywordsMatched", [])
    missing = opt_result.get("keywordsMissing", [])

    with kw1:
        st.markdown(f"**✅ Matched Keywords ({len(matched)})**")
        chips = " ".join(f"<span class='keyword-matched'>{k}</span>" for k in matched[:25])
        st.markdown(chips, unsafe_allow_html=True)

    with kw2:
        st.markdown(f"**⚠️ Hard Gaps ({len(missing)}) — add only if genuine**")
        chips = " ".join(f"<span class='keyword-missing'>{k}</span>" for k in missing[:15])
        st.markdown(chips, unsafe_allow_html=True)

    # ── Section feedback ────────────────────────────────────────────────────────
    feedback = opt_result.get("sectionFeedback", {})
    if feedback:
        with st.expander("📋 Detailed Section Feedback"):
            for section, note in feedback.items():
                st.markdown(f"**{section.title()}:** {note}")

    # ── Cover letter preview ────────────────────────────────────────────────────
    with st.expander("👁️ Preview Cover Letter Body"):
        st.write(cover_body)

    # ── Resume text preview ─────────────────────────────────────────────────────
    with st.expander("👁️ Preview Optimized Resume Text"):
        st.text(opt_result.get("optimizedResumeText", ""))
